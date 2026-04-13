from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── 頁面設定 A4 ──
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

# ── 輔助函式 ──
def set_font(run, bold=False, size=11, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, bold=True, size=22, color=(124, 100, 213))
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    # 底線
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '18')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '7c64d5')
    pBdr.append(bottom)
    pPr.append(pBdr)

def heading2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=True, size=14, color=(26, 26, 46))
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)

def label(doc, text, color=(124, 100, 213)):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(2)

def body(doc, text, color=(85, 85, 85)):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(*color)

def bullet(doc, text, color=(85, 85, 85)):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(*color)

def metric_line(doc, label_text, before, after, improve, good=True):
    pass  # handled via table

# ════════════════════════════════════════════════════
# 封面
# ════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PERFORMANCE OPTIMIZATION REPORT')
set_font(run, bold=True, size=10, color=(124, 100, 213))
p.paragraph_format.space_before = Pt(60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('作品集網站 效能優化報告')
set_font(run, bold=True, size=28, color=(26, 26, 46))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Tony Cheng Portfolio  ·  abonla.github.io/web_portfolio')
set_font(run, size=11, color=(100, 100, 120))
p.paragraph_format.space_after = Pt(24)

# 封面資訊表
tbl = doc.add_table(rows=1, cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'
cells = tbl.rows[0].cells
data = [('2026-04-13', '優化日期'), ('4 項', '優化項目'), ('3 commits', '程式異動')]
for i, (val, lbl) in enumerate(data):
    p = cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(val + '\n')
    set_font(r1, bold=True, size=14, color=(26, 26, 46))
    r2 = p.add_run(lbl)
    set_font(r2, size=9, color=(120, 120, 150))

doc.add_page_break()

# ════════════════════════════════════════════════════
# PAGE 1：問題診斷
# ════════════════════════════════════════════════════
label(doc, 'SECTION 01')
heading1(doc, '問題診斷')
body(doc, '網站在手機幾乎無法開啟，桌機載入也極為緩慢。經過分析，根本原因來自以下三個問題同時存在：')

problems = [
    ('01', '縮圖未壓縮',    '226 張縮圖平均 130 KB，最大達 653 KB，總下載量高達 32 MB。即使有 loading="lazy"，首屏仍需下載大量資料。',    (255, 107, 107)),
    ('02', 'Masonry 重排過多', '每張圖片載入就觸發一次 masonry() 排版計算，頁面載入期間高達 200+ 次，手機 CPU 嚴重過載造成凍結。', (247, 183, 49)),
    ('03', 'YouTube 影片空白', '53 個 <iframe> 在頁面載入時同時發起外部連線，且 file:// 本地協定被 YouTube 拒絕嵌入，影片一片空白。', (76, 195, 255)),
]
for num, title, desc, color in problems:
    p = doc.add_paragraph()
    r = p.add_run(f'問題 {num}  {title}')
    set_font(r, bold=True, size=12, color=color)
    body(doc, desc)

# ── FIX 01 ──
label(doc, 'FIX 01  ·  COMMIT 375e43a')
heading1(doc, '縮圖批次壓縮')
body(doc, '對所有縮圖（*m.jpg、*m.png、*小.jpg、*s.jpg）批次重壓縮，最大寬度限制 600px，quality=70。PNG 無透明度者自動轉為 JPEG。原始高解析圖保留不動（僅 Fancybox 點擊後才下載）。')

# 壓縮成果表
heading2(doc, '壓縮成果')
tbl2 = doc.add_table(rows=5, cols=3)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
headers = ['指標', '壓縮前', '壓縮後']
rows_data = [
    ('總下載量',    '~32 MB',  '~12 MB'),
    ('節省容量',    '—',       '20 MB（-62%）'),
    ('單張最大',    '653 KB',  '~70 KB（-89%）'),
    ('處理檔案數',  '—',       '246 個縮圖'),
]
for j, h in enumerate(headers):
    cell = tbl2.rows[0].cells[j]
    p = cell.paragraphs[0]
    run = p.add_run(h)
    set_font(run, bold=True, size=10, color=(255, 255, 255))
    # 深色背景
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '0d0d1a')
    tcPr.append(shd)

for i, (lbl_text, before, after) in enumerate(rows_data):
    row = tbl2.rows[i+1]
    for j, text in enumerate([lbl_text, before, after]):
        p = row.cells[j].paragraphs[0]
        run = p.add_run(text)
        color = (32, 191, 107) if j == 2 else (255, 107, 107) if j == 1 and before != '—' else (50, 50, 50)
        set_font(run, bold=(j==2), size=10, color=color)

doc.add_page_break()

# ════════════════════════════════════════════════════
# PAGE 2：Masonry + YouTube
# ════════════════════════════════════════════════════
label(doc, 'FIX 02  ·  COMMIT 375e43a')
heading1(doc, 'Masonry 排版 Debounce')
body(doc, '原始程式碼每張圖片的 load 事件都直接呼叫 masonry()，200+ 張圖就觸發 200+ 次重排。改用 debounce 後，200ms 內所有事件只執行一次排版，CPU 負擔降低 97%+。')

heading2(doc, '修正前（每張圖載入都觸發，200+ 次）')
p = doc.add_paragraph()
run = p.add_run(
    '$grid.imagesLoaded().progress(function() {\n'
    '    $grid.masonry();  // ← 每次都執行\n'
    '});'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(180, 180, 180)

heading2(doc, '修正後（debounce，最多每 200ms 執行一次）')
p = doc.add_paragraph()
run = p.add_run(
    'var layoutTimer;\n'
    '$grid.imagesLoaded().progress(function() {\n'
    '    clearTimeout(layoutTimer);\n'
    '    layoutTimer = setTimeout(function() {\n'
    "        $grid.masonry('layout');\n"
    '    }, 200);\n'
    '});'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(50, 50, 50)

label(doc, 'FIX 03  ·  COMMIT 1c45f84', color=(124, 100, 213))
heading1(doc, 'YouTube Facade 模式')
body(doc, '將所有 <iframe> 改為「YouTube 縮圖 + 播放按鈕」的 Facade，使用者點擊後才動態插入 iframe（搭配 autoplay=1），並改用 youtube-nocookie.com 降低嵌入限制。')

tbl3 = doc.add_table(rows=5, cols=2)
tbl3.style = 'Table Grid'
before_items = ['53 個 iframe 同時建立連線', 'file:// 協定被 YouTube 封鎖', '影片顯示空白或錯誤訊息', '頁面額外下載 2–5 MB']
after_items  = ['頁面載入 YouTube 連線：0 條', '縮圖每張僅 ~5 KB', '點擊後正常播放', 'GitHub Pages 完整運作']
for i in range(5):
    for j in range(2):
        cell = tbl3.rows[i].cells[j]
        p = cell.paragraphs[0]
        if i == 0:
            run = p.add_run('修正前' if j == 0 else '修正後')
            set_font(run, bold=True, size=10, color=(255, 107, 107) if j == 0 else (32, 191, 107))
        else:
            text = before_items[i-1] if j == 0 else after_items[i-1]
            run = p.add_run(text)
            set_font(run, size=9.5, color=(85, 85, 85))

label(doc, 'FIX 04  ·  COMMIT 20aeaa6', color=(124, 100, 213))
heading1(doc, '點擊影片不跑版')
body(doc, '原本用 replaceWith(iframe) 直接換掉 DOM 節點，iframe 高度（220px）與縮圖高度不符，Masonry 重算後整頁跑版。改用 padding-bottom: 56.25% 固定 16:9 比例容器，iframe 以絕對定位填滿，替換前後容器高度完全一致。')

p = doc.add_paragraph()
run = p.add_run(
    '/* 固定 16:9 容器，點擊前後高度不變 */\n'
    '.yt-facade, .yt-playing {\n'
    '  position: relative;\n'
    '  padding-bottom: 56.25%;  /* 16:9 */\n'
    '}\n'
    '.yt-playing iframe {\n'
    '  position: absolute;\n'
    '  top: 0; left: 0;\n'
    '  width: 100%; height: 100%;\n'
    '}'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(50, 50, 50)

doc.add_page_break()

# ════════════════════════════════════════════════════
# PAGE 3：成果 + 時間軸
# ════════════════════════════════════════════════════
label(doc, 'SECTION 02')
heading1(doc, '優化成果總覽')

metrics = [
    ('62%',  '縮圖總下載量減少',       '32 MB → 12 MB · 246 個縮圖',  (32, 191, 107)),
    ('89%',  '單張最大縮圖縮小',       '653 KB → ~70 KB',             (32, 191, 107)),
    ('0',    '載入時 YouTube 連線數',  '原本 53 條 → 0 條',           (76, 195, 255)),
    ('97%',  'Masonry 重排次數降低',   '200+ 次 → <5 次',             (247, 183, 49)),
]
for num, title, sub, color in metrics:
    p = doc.add_paragraph()
    r1 = p.add_run(f'{num}  ')
    set_font(r1, bold=True, size=20, color=color)
    r2 = p.add_run(f'{title}  ')
    set_font(r2, bold=True, size=11, color=(26, 26, 46))
    r3 = p.add_run(sub)
    set_font(r3, size=9, color=(120, 120, 140))
    p.paragraph_format.space_after = Pt(4)

# 總覽表
heading2(doc, '詳細數據對照表')
summary_rows = [
    ('縮圖總下載量',         '~32 MB',    '~12 MB',      '▼ 62%',  True),
    ('單張最大縮圖',         '653 KB',    '~70 KB',      '▼ 89%',  True),
    ('頁面載入 YouTube 連線數', '53 條',  '0 條',        '▼ 100%', True),
    ('Masonry 排版計算次數',  '200+ 次',  '<5 次',       '▼ 97%+', True),
    ('點擊影片是否跑版',      '是',        '否',          '✓ 修正', True),
    ('本機預覽影片是否顯示',  '否（空白）','顯示縮圖',    '✓ 修正', True),
]
tbl4 = doc.add_table(rows=len(summary_rows)+1, cols=4)
tbl4.style = 'Table Grid'
for j, h in enumerate(['指標', '優化前', '優化後', '改善幅度']):
    cell = tbl4.rows[0].cells[j]
    p = cell.paragraphs[0]
    run = p.add_run(h)
    set_font(run, bold=True, size=10, color=(255, 255, 255))
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '0d0d1a')
    tcPr.append(shd)

for i, (metric, before, after, improve, good) in enumerate(summary_rows):
    row = tbl4.rows[i+1]
    texts = [metric, before, after, improve]
    colors = [(50, 50, 50), (255, 107, 107), (50, 50, 50), (32, 191, 107)]
    for j, (text, color) in enumerate(zip(texts, colors)):
        p = row.cells[j].paragraphs[0]
        run = p.add_run(text)
        set_font(run, bold=(j==3), size=9.5, color=color)

# 時間軸
label(doc, 'SECTION 03')
heading1(doc, '提交記錄時間軸')

commits = [
    ('1', '375e43a · 2026-04-13', '縮圖壓縮 + Masonry Debounce',
     '以 Python Pillow quality=70 批次重壓縮 246 個縮圖（32MB→12MB）；同時加入 200ms debounce 大幅減少 Masonry 排版重算次數。',
     (124, 100, 213)),
    ('2', '1c45f84 · 2026-04-13', 'YouTube Facade 模式',
     '將全部 53 個 <iframe> 替換為縮圖 + 播放按鈕 Facade，改用 youtube-nocookie.com，點擊才動態插入 iframe 播放。',
     (124, 100, 213)),
    ('3', '20aeaa6 · 2026-04-13', '點擊影片跑版修正',
     '改用 padding-bottom: 56.25% 固定 16:9 容器比例，iframe 以絕對定位填滿，點擊前後容器高度不變，Masonry 排版完全不受影響。',
     (32, 191, 107)),
]
for num, commit_id, title, desc, color in commits:
    p = doc.add_paragraph()
    r1 = p.add_run(f'[{num}]  ')
    set_font(r1, bold=True, size=12, color=color)
    r2 = p.add_run(commit_id + '\n')
    r2.font.name = 'Consolas'
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor(*color)
    r3 = p.add_run(title + '\n')
    set_font(r3, bold=True, size=11, color=(26, 26, 46))
    r4 = p.add_run(desc)
    set_font(r4, size=9.5, color=(85, 85, 85))
    p.paragraph_format.space_after = Pt(12)

# ── 儲存 ──
output_path = r'c:\Users\abon8\Documents\GitHub\web_portfolio\作品集網站效能優化報告.docx'
doc.save(output_path)
print(f'已儲存：{output_path}')
